import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets inner margins (padding) for a table cell (in twentieths of a point, dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    """Applies a clean light-grey border style to a table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="cbd5e0"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="cbd5e0"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="cbd5e0"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="cbd5e0"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="cbd5e0"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="cbd5e0"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def clean_inline_formatting(text):
    """Extracts text by removing HTML bold/italic/code tags and markdown bold/italic markers."""
    text = re.sub(r'</?(b|strong|i|em|code|span|div)[^>]*>', '', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = text.replace('&nbsp;', ' ')
    text = text.replace('&rarr;', '➔')
    return text

def parse_runs(paragraph, line_text):
    """Parses text line and adds formatted runs (bold, italic, code) to the paragraph."""
    # Pattern to find bold, italic, code, or plain text
    # We support simple markdown formatting: **bold**, *italic*, `code`
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|&nbsp;|&rarr;)', line_text)
    
    for token in tokens:
        if not token:
            continue
        
        # Check token type
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith('`') and token.endswith('`'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
            # Add a light grey background to code run if possible
            rPr = run._r.get_or_add_rPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="f7fafc"/>')
            rPr.append(shd)
        elif token == '&nbsp;':
            paragraph.add_run(' ')
        elif token == '&rarr;':
            paragraph.add_run(' ➔ ')
        else:
            # Handle plain text, check for inline HTML tags like <b> or <i>
            sub_tokens = re.split(r'(<b>.*?</b>|<i>.*?</i>|<code>.*?</code>)', token)
            for sub_token in sub_tokens:
                if not sub_token:
                    continue
                if sub_token.startswith('<b>') and sub_token.endswith('</b>'):
                    run = paragraph.add_run(sub_token[3:-4])
                    run.bold = True
                elif sub_token.startswith('<i>') and sub_token.endswith('</i>'):
                    run = paragraph.add_run(sub_token[3:-4])
                    run.italic = True
                elif sub_token.startswith('<code>') and sub_token.endswith('</code>'):
                    run = paragraph.add_run(sub_token[6:-7])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9.5)
                else:
                    paragraph.add_run(sub_token)

def create_docx(md_path: str, docx_path: str):
    doc = Document()
    
    # Page setup - A4, 1 inch (2.54 cm) margins
    for section in doc.sections:
        section.page_width = Inches(8.27)  # A4 Width
        section.page_height = Inches(11.69)  # A4 Height
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    styles = doc.styles
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    normal_style.paragraph_format.line_spacing = 1.25
    normal_style.paragraph_format.space_after = Pt(6)
    
    # ------------------ COVER PAGE ------------------
    # We construct the cover page programmatically to match the official BITS template
    
    # Add title spacing
    for _ in range(2):
        doc.add_paragraph()
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("EFFICIENCY-OPTIMIZED GENERATIVE PARADIGMS FOR\nLARGE-SCALE BIOMOLECULAR DESIGN")
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(16)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(0x0c, 0x23, 0x40)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run("BITS ZG628T: Dissertation")
    run_sub.font.name = 'Times New Roman'
    run_sub.font.size = Pt(12)
    run_sub.bold = True
    
    p_by = doc.add_paragraph()
    p_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_by.paragraph_format.space_after = Pt(6)
    run_by = p_by.add_run("by")
    run_by.font.name = 'Times New Roman'
    run_by.font.size = Pt(11)
    run_by.italic = True
    
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_author.paragraph_format.space_after = Pt(4)
    run_author = p_author.add_run("Akik Jana")
    run_author.font.name = 'Times New Roman'
    run_author.font.size = Pt(13)
    run_author.bold = True
    
    p_id = doc.add_paragraph()
    p_id.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_id.paragraph_format.space_after = Pt(36)
    run_id = p_id.add_run("(2024AB05287)")
    run_id.font.name = 'Times New Roman'
    run_id.font.size = Pt(11)
    
    p_carried = doc.add_paragraph()
    p_carried.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_carried.paragraph_format.space_after = Pt(6)
    run_carried = p_carried.add_run("Dissertation work carried out at")
    run_carried.font.name = 'Times New Roman'
    run_carried.font.size = Pt(11)
    run_carried.italic = True
    
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_org.paragraph_format.space_after = Pt(36)
    run_org = p_org.add_run("Mu Sigma, Bangalore")
    run_org.font.name = 'Times New Roman'
    run_org.font.size = Pt(12)
    run_org.bold = True
    
    p_submitted = doc.add_paragraph()
    p_submitted.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_submitted.paragraph_format.left_indent = Inches(0.75)
    p_submitted.paragraph_format.right_indent = Inches(0.75)
    p_submitted.paragraph_format.space_after = Pt(36)
    run_submitted = p_submitted.add_run("Submitted in partial fulfilment of M.Tech. in Artificial Intelligence and Machine Learning degree programme")
    run_submitted.font.name = 'Times New Roman'
    run_submitted.font.size = Pt(11)
    run_submitted.italic = True
    
    p_supervision = doc.add_paragraph()
    p_supervision.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_supervision.paragraph_format.space_after = Pt(6)
    run_super = p_supervision.add_run("Under the Supervision of")
    run_super.font.name = 'Times New Roman'
    run_super.font.size = Pt(11)
    run_super.italic = True
    
    p_supervisor = doc.add_paragraph()
    p_supervisor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_supervisor.paragraph_format.space_after = Pt(4)
    run_supervisor = p_supervisor.add_run("Dr. Arnab Bandyopadhyay")
    run_supervisor.font.name = 'Times New Roman'
    run_supervisor.font.size = Pt(12)
    run_supervisor.bold = True
    
    p_super_org = doc.add_paragraph()
    p_super_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_super_org.paragraph_format.space_after = Pt(40)
    run_super_org = p_super_org.add_run("RnD Division, Dr. Reddy's Laboratories, Hyderabad")
    run_super_org.font.name = 'Times New Roman'
    run_super_org.font.size = Pt(11)
    
    # BITS LOGO
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_after = Pt(12)
    logo_path = "/Users/akikjana/Documents/BiomolecularDesign/src/bits_logo.png"
    if os.path.exists(logo_path):
        p_logo.add_run().add_picture(logo_path, width=Inches(1.2))
    else:
        p_logo.add_run("[BITS Pilani Logo]")
        
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_inst.paragraph_format.space_after = Pt(4)
    run_inst = p_inst.add_run("BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE")
    run_inst.font.name = 'Times New Roman'
    run_inst.font.size = Pt(13)
    run_inst.bold = True
    run_inst.font.color.rgb = RGBColor(0x0c, 0x23, 0x40)
    
    p_loc = doc.add_paragraph()
    p_loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_loc.paragraph_format.space_after = Pt(12)
    run_loc = p_loc.add_run("PILANI (RAJASTHAN)")
    run_loc.font.name = 'Times New Roman'
    run_loc.font.size = Pt(12)
    run_loc.bold = True
    run_loc.font.color.rgb = RGBColor(0x0c, 0x23, 0x40)
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run("June 2026")
    run_date.font.name = 'Times New Roman'
    run_date.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ------------------ PARSE BODY CONTENT ------------------
    with open(md_path, 'r', encoding='utf-8') as f:
        md_lines = f.readlines()
        
    # We skip everything before the first "---" (which corresponds to the cover info)
    skip_mode = True
    in_code_block = False
    code_content = []
    
    in_table = False
    table_rows = []
    
    in_flowchart = False
    flowchart_steps = []
    
    in_sig_block = False
    
    i = 0
    while i < len(md_lines):
        line = md_lines[i].strip()
        
        # Skip cover page info from MD (since we generated a nice one above)
        if skip_mode:
            if line == '---':
                skip_mode = False
            i += 1
            continue
            
        # Detect Page Breaks
        if 'class="page-break"' in line or 'page-break-before' in line:
            doc.add_page_break()
            i += 1
            continue
            
        # Detect Signature Block HTML
        if '<div class="signature-block">' in line:
            in_sig_block = True
            i += 1
            continue
            
        if in_sig_block:
            if '</div>' in line and not ('<div' in line or '<span' in line or '<p' in line):
                # Close signature block and render
                in_sig_block = False
                
                # Render the 2-column signature table
                sig_table = doc.add_table(rows=5, cols=2)
                sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Student column (Col 0)
                cell_std = sig_table.cell(0, 0)
                p_std_line = cell_std.paragraphs[0]
                p_std_line.paragraph_format.space_after = Pt(4)
                p_std_line.add_run("_______________________________\n").bold = True
                run_st = p_std_line.add_run("Signature of the Student\n")
                run_st.bold = True
                run_st.font.size = Pt(11)
                
                p_std_details = cell_std.add_paragraph()
                p_std_details.paragraph_format.space_after = Pt(2)
                p_std_details.add_run("Name: ").bold = True
                p_std_details.add_run("Akik Jana\n")
                p_std_details.add_run("Date: ").bold = True
                p_std_details.add_run("June 16, 2026\n")
                p_std_details.add_run("Place: ").bold = True
                p_std_details.add_run("Bangalore")
                
                # Supervisor column (Col 1)
                cell_sup = sig_table.cell(0, 1)
                p_sup_line = cell_sup.paragraphs[0]
                p_sup_line.paragraph_format.space_after = Pt(4)
                p_sup_line.add_run("_______________________________\n").bold = True
                run_spt = p_sup_line.add_run("Signature of the Supervisor\n")
                run_spt.bold = True
                run_spt.font.size = Pt(11)
                
                p_sup_details = cell_sup.add_paragraph()
                p_sup_details.paragraph_format.space_after = Pt(2)
                p_sup_details.add_run("Name: ").bold = True
                p_sup_details.add_run("Dr. Arnab Bandyopadhyay\n")
                p_sup_details.add_run("Date: ").bold = True
                p_sup_details.add_run("June 16, 2026\n")
                p_sup_details.add_run("Place: ").bold = True
                p_sup_details.add_run("Hyderabad")
                
                # Add spacing after signature block
                doc.add_paragraph()
            i += 1
            continue
            
        # Detect Flowchart Block HTML
        if '<div class="flowchart-container">' in line:
            in_flowchart = True
            flowchart_steps = []
            i += 1
            continue
            
        if in_flowchart:
            # Parse flowchart box info
            if '<div class="flowchart-box' in line:
                # Find next title and description
                box_title = ""
                box_desc = ""
                while True:
                    i += 1
                    sub_line = md_lines[i].strip()
                    if '<h3>' in sub_line:
                        box_title = re.sub(r'</?h3[^>]*>', '', sub_line)
                    elif '<p>' in sub_line:
                        box_desc = re.sub(r'</?p[^>]*>', '', sub_line)
                    elif '</div>' in sub_line:
                        break
                flowchart_steps.append((box_title, box_desc))
            elif '</div>' in line and not ('<div' in line or '<p' in line or '<h3>' in line):
                in_flowchart = False
                
                # Render Flowchart as a beautiful 1x5 table
                flow_table = doc.add_table(rows=1, cols=5)
                flow_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Set specific column widths
                widths = [Inches(2.0), Inches(0.3), Inches(2.0), Inches(0.3), Inches(2.0)]
                for row in flow_table.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width
                
                # Col 0: Step 1
                c0 = flow_table.cell(0, 0)
                set_cell_background(c0, "ebf8ff") # Light blue
                set_cell_margins(c0, 100, 100, 100, 100)
                p0 = c0.paragraphs[0]
                p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_p0_t = p0.add_run(flowchart_steps[0][0] + "\n")
                run_p0_t.bold = True
                run_p0_t.font.size = Pt(10)
                run_p0_t.font.color.rgb = RGBColor(0x0c, 0x23, 0x40)
                run_p0_d = p0.add_run(flowchart_steps[0][1])
                run_p0_d.font.size = Pt(8.5)
                
                # Col 1: Arrow
                c1 = flow_table.cell(0, 1)
                set_cell_margins(c1, 100, 100, 50, 50)
                p1 = c1.paragraphs[0]
                p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_p1 = p1.add_run("\n➔")
                run_p1.bold = True
                run_p1.font.size = Pt(14)
                run_p1.font.color.rgb = RGBColor(0xab, 0xbd, 0xd0)
                
                # Col 2: Step 2
                c2 = flow_table.cell(0, 2)
                set_cell_background(c2, "f0fff4") # Light green
                set_cell_margins(c2, 100, 100, 100, 100)
                p2 = c2.paragraphs[0]
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_p2_t = p2.add_run(flowchart_steps[1][0] + "\n")
                run_p2_t.bold = True
                run_p2_t.font.size = Pt(10)
                run_p2_t.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
                run_p2_d = p2.add_run(flowchart_steps[1][1])
                run_p2_d.font.size = Pt(8.5)
                
                # Col 3: Arrow
                c3 = flow_table.cell(0, 3)
                set_cell_margins(c3, 100, 100, 50, 50)
                p3 = c3.paragraphs[0]
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_p3 = p3.add_run("\n➔")
                run_p3.bold = True
                run_p3.font.size = Pt(14)
                run_p3.font.color.rgb = RGBColor(0xab, 0xbd, 0xd0)
                
                # Col 4: Step 3
                c4 = flow_table.cell(0, 4)
                set_cell_background(c4, "fffaf0") # Light orange
                set_cell_margins(c4, 100, 100, 100, 100)
                p4 = c4.paragraphs[0]
                p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_p4_t = p4.add_run(flowchart_steps[2][0] + "\n")
                run_p4_t.bold = True
                run_p4_t.font.size = Pt(10)
                run_p4_t.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)
                run_p4_d = p4.add_run(flowchart_steps[2][1])
                run_p4_d.font.size = Pt(8.5)
                
                doc.add_paragraph() # spacing
            i += 1
            continue

        # Code Block detection
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_content = []
            else:
                in_code_block = False
                # Render code block in a light-grey callout box
                tbl = doc.add_table(rows=1, cols=1)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = tbl.cell(0, 0)
                set_cell_background(cell, "f7fafc")
                set_cell_margins(cell, 150, 150, 200, 200)
                
                p = cell.paragraphs[0]
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(0)
                code_text = "\n".join(code_content)
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)
                
                # Style border
                tblPr = tbl._tbl.tblPr
                borders = parse_xml(
                    f'<w:tblBorders {nsdecls("w")}>'
                    f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="e2e8f0"/>'
                    f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="e2e8f0"/>'
                    f'  <w:left w:val="single" w:sz="8" w:space="0" w:color="3182ce"/>' # Left accent border blue
                    f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="e2e8f0"/>'
                    f'</w:tblBorders>'
                )
                tblPr.append(borders)
                
                doc.add_paragraph() # spacer
            i += 1
            continue
            
        if in_code_block:
            code_content.append(md_lines[i].rstrip('\n'))
            i += 1
            continue
            
        # Table detection
        if line.startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            
            # Skip separator line (e.g. | :--- | :--- |)
            if not re.match(r'^\|\s*[:\-|\s]+\s*\|$', line):
                # Parse columns
                cols = [c.strip() for c in line.split('|')[1:-1]]
                table_rows.append(cols)
                
            i += 1
            continue
        else:
            if in_table:
                # End of table, render it
                in_table = False
                if len(table_rows) > 0:
                    num_cols = len(table_rows[0])
                    doc_table = doc.add_table(rows=len(table_rows), cols=num_cols)
                    doc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    set_table_borders(doc_table)
                    
                    # Style and populate
                    for r_idx, row_data in enumerate(table_rows):
                        row = doc_table.rows[r_idx]
                        for c_idx, cell_text in enumerate(row_data):
                            if c_idx >= num_cols:
                                break
                            cell = row.cells[c_idx]
                            set_cell_margins(cell, 100, 100, 120, 120)
                            
                            # Header styling
                            if r_idx == 0:
                                set_cell_background(cell, "ebf8ff") # Light blue header
                                p = cell.paragraphs[0]
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = p.add_run(clean_inline_formatting(cell_text))
                                run.bold = True
                                run.font.size = Pt(9.5)
                                run.font.color.rgb = RGBColor(0x2b, 0x6c, 0xb0)
                            else:
                                # Alternating row colors
                                if r_idx % 2 == 0:
                                    set_cell_background(cell, "f7fafc")
                                p = cell.paragraphs[0]
                                run = p.add_run(clean_inline_formatting(cell_text))
                                run.font.size = Pt(9.0)
                    
                    doc.add_paragraph() # spacer
            # continue parsing current line
            
        # Detect Images
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)$', line)
        if img_match:
            caption = img_match.group(1)
            img_path = img_match.group(2)
            
            # Embed image in Word
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(12)
            p_img.paragraph_format.space_after = Pt(4)
            
            if os.path.exists(img_path):
                # Adjust width based on type
                if "rendering" in img_path:
                    p_img.add_run().add_picture(img_path, width=Inches(3.8))
                else:
                    p_img.add_run().add_picture(img_path, width=Inches(4.5))
            else:
                p_img.add_run(f"[Image missing: {os.path.basename(img_path)}]")
                
            # Add Caption
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(12)
            run_cap = p_cap.add_run(f"Figure: {caption}")
            run_cap.italic = True
            run_cap.font.size = Pt(9.5)
            run_cap.font.color.rgb = RGBColor(0x4a, 0x55, 0x68)
            
            i += 1
            continue
            
        # Headings
        if line.startswith('# '):
            p = doc.add_paragraph(style='Heading 1')
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[2:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(18)
            run.bold = True
            run.font.color.rgb = RGBColor(0x0c, 0x23, 0x40)
            i += 1
            continue
            
        if line.startswith('## '):
            p = doc.add_paragraph(style='Heading 2')
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[3:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
            i += 1
            continue
            
        if line.startswith('### '):
            p = doc.add_paragraph(style='Heading 3')
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[4:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11.5)
            run.bold = True
            run.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)
            i += 1
            continue
            
        if line.startswith('#### '):
            p = doc.add_paragraph(style='Heading 4')
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[5:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)
            i += 1
            continue
            
        # Lists (bulleted)
        if line.startswith('* ') or line.startswith('- '):
            # Trim prefix
            bullet_text = line[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            parse_runs(p, bullet_text)
            i += 1
            continue
            
        # Empty Line
        if not line:
            # We don't add blank paragraphs for every blank line to maintain proper Word spacing,
            # unless it's a specific divider or we want vertical spacing.
            i += 1
            continue
            
        # Regular Paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        parse_runs(p, line)
        
        i += 1
        
    # Render any remaining table at the end of the file (e.g., Abbreviations table)
    if in_table and len(table_rows) > 0:
        num_cols = len(table_rows[0])
        doc_table = doc.add_table(rows=len(table_rows), cols=num_cols)
        set_table_borders(doc_table)
        for r_idx, row_data in enumerate(table_rows):
            row = doc_table.rows[r_idx]
            for c_idx, cell_text in enumerate(row_data):
                if c_idx >= num_cols:
                    break
                cell = row.cells[c_idx]
                set_cell_margins(cell, 100, 100, 120, 120)
                if r_idx == 0:
                    set_cell_background(cell, "ebf8ff")
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(clean_inline_formatting(cell_text))
                    run.bold = True
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0x2b, 0x6c, 0xb0)
                else:
                    if r_idx % 2 == 0:
                        set_cell_background(cell, "f7fafc")
                    p = cell.paragraphs[0]
                    run = p.add_run(clean_inline_formatting(cell_text))
                    run.font.size = Pt(9.0)
                    
    doc.save(docx_path)
    print(f"[DOCX] Successfully generated DOCX report at: {docx_path}")

if __name__ == "__main__":
    brain_dir = "/Users/akikjana/.gemini/antigravity-cli/brain/4cdb7261-e55b-4efc-9ffa-c6509d76c9c2"
    md_file = os.path.join(brain_dir, "mid_semester_report.md")
    docx_file = "/Users/akikjana/Documents/BiomolecularDesign/mid_semester_report.docx"
    
    create_docx(md_file, docx_file)
